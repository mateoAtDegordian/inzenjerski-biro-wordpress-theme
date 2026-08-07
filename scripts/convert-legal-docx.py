#!/usr/bin/env python3
"""Convert the approved legal DOCX files into semantic theme HTML fragments."""

from __future__ import annotations

import html
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def iter_blocks(parent: DocumentType):
    for child in parent.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def inline(paragraph: Paragraph) -> str:
    parts: list[str] = []
    runs: list[Run] = []
    for child in paragraph._p.iterchildren():
        if child.tag.endswith("}r"):
            runs.append(Run(child, paragraph))
        elif child.tag.endswith("}hyperlink"):
            runs.extend(Run(item, paragraph) for item in child.iterchildren() if item.tag.endswith("}r"))

    for run in runs:
        value = html.escape(run.text).replace("\n", "<br>")
        if not value:
            continue
        if run.bold:
            value = f"<strong>{value}</strong>"
        if run.italic:
            value = f"<em>{value}</em>"
        if run.underline:
            value = f"<u>{value}</u>"
        parts.append(value)
    return "".join(parts).strip()


def table_html(table: Table) -> str:
    if len(table.rows) == 1 and len(table.columns) == 1:
        contents = []
        for p in table.cell(0, 0).paragraphs:
            value = inline(p)
            if value:
                contents.append(f"<p>{value}</p>")
        return '<aside class="legal-callout">' + "".join(contents) + "</aside>"

    rows = []
    for index, row in enumerate(table.rows):
        tag = "th" if index == 0 else "td"
        cells = []
        for cell in row.cells:
            contents = "<br>".join(inline(p) for p in cell.paragraphs if inline(p))
            cells.append(f"<{tag}>{contents}</{tag}>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return '<div class="legal-table-wrap"><table>' + "".join(rows) + "</table></div>"


def convert(source: Path, destination: Path) -> None:
    document = Document(source)
    output: list[str] = []
    in_list = False
    first_text = True

    for block in iter_blocks(document):
        if isinstance(block, Table):
            if in_list:
                output.append("</ul>")
                in_list = False
            output.append(table_html(block))
            continue

        value = inline(block)
        if not value:
            continue
        style = block.style.name if block.style else "Normal"
        is_list = style == "List Paragraph"
        if is_list and not in_list:
            output.append("<ul>")
            in_list = True
        elif not is_list and in_list:
            output.append("</ul>")
            in_list = False

        if is_list:
            output.append(f"<li>{value}</li>")
        elif first_text:
            output.append(f"<h1>{value}</h1>")
            first_text = False
        elif style == "Heading 1":
            output.append(f"<h2>{value}</h2>")
        elif style == "Heading 2":
            output.append(f"<h3>{value}</h3>")
        else:
            output.append(f"<p>{value}</p>")

    if in_list:
        output.append("</ul>")

    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text("\n".join(output) + "\n", encoding="utf-8")


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: convert-legal-docx.py INPUT_DIR OUTPUT_DIR", file=sys.stderr)
        return 2
    source_dir = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])
    mapping = {
        "Politika-privatnosti-INGBIRO-2026.docx": "privacy-hr.html",
        "Privacy-Policy-INGBIRO-2026-EN.docx": "privacy-en.html",
        "Politika_koris%CC%8Ctenja_kolac%CC%8Cic%CC%81a.docx": "cookies-hr.html",
        "Cookie-Policy-INGBIRO-EN.docx": "cookies-en.html",
    }
    for source_name, destination_name in mapping.items():
        convert(source_dir / source_name, output_dir / destination_name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
